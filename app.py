import yt_dlp
import assemblyai as aai
import os
import tempfile
import shutil
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from collections import Counter
import google.generativeai as genai
import json
from dotenv import load_dotenv
from docx import Document
import re
import requests
import logging
from gtts import gTTS
from pydub import AudioSegment
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from flask import send_from_directory

app = Flask(__name__)
CORS(app)

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), 'output')
# os.makedirs(OUTPUT_DIR, exist_ok=True)


TEMP_DIR = os.path.join(os.path.dirname(__file__), 'Temp')
# os.makedirs(TEMP_DIR, exist_ok=True)

# Import the model configuration
from model_config import model

# Configure logging
logging.basicConfig(level=logging.INFO)

# Load environment variables
load_dotenv()
ASSEMBLYAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# Configure APIs
aai.settings.api_key = ASSEMBLYAI_API_KEY
genai.configure(api_key=GOOGLE_API_KEY)
# model = genai.GenerativeModel("gemini-1.5-flash")   # gentai model name (Must be free version)

import subprocess
# Transcribe video using AssemblyAI and save to result (entry point 2)
def transcribe_video(url):
    transcriber = aai.Transcriber()  
    # temp_dir = tempfile.mkdtemp()  
    audio_file = os.path.join(TEMP_DIR, "audio.mp3") 

    if os.environ.get("RENDER", "0") == "1":
        ffmpeg_path = shutil.which("ffmpeg") or "/usr/bin/ffmpeg"
    else:
        # On local Windows machine, use your local ffmpeg path
        ffmpeg_path = "C:\\ffmpeg\\ffmpeg.exe"

    # ffmpeg_path = "C:\\ffmpeg\\ffmpeg.exe"  # Path to ffmpeg in local machine
    if not ffmpeg_path:
        return "Error: ffmpeg not found in PATH. Please install it or specify the path manually."
    

    abc_file = "index.txt"
    abc = os.getenv("ABCYT", "").replace("\\n", "\n")

    # Write into Netscape-format file only if abc has content
    if abc.strip():
        with open(abc_file, "w", encoding="utf-8") as f:
            f.write("# Netscape HTTP Cookie File\n")
            for item in abc.split("; "):
                if "=" in item:
                    name, value = item.split("=", 1)
                    f.write(f".youtube.com\tTRUE\t/\tFALSE\t0\t{name}\t{value}\n")

    # yt-dlp options
    ydl_opts = {
        'format': 'bestaudio/best',
        'outtmpl': os.path.join(TEMP_DIR, "audio"),  # No extension
        'quiet': True,
        'ffmpeg_location': ffmpeg_path,
        'cookiefile': abc_file if abc.strip() else None,
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'mp3',
        }],
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])                 #Downloades the audio in mp3 format
    except yt_dlp.DownloadError as e:
        return f"Download Error: {str(e)}"
    finally:        
        if os.path.exists(abc_file):
            try:
                os.remove(abc_file)
            except Exception:
                pass

    if not os.path.exists(audio_file):
        return f"File Error: {audio_file} was not created."

    try:
        transcript = transcriber.transcribe(audio_file)     # Transcribes the audio file using AssemblyAI
        if transcript.status == aai.TranscriptStatus.error:
            return f"Transcription Error: {transcript.error}"
        result = transcript.text
        os.remove(audio_file)
        print(f"Removed audio file: {audio_file}") # let's see if it works (Removes temp audio file (audio))
        # os.rmdir(temp_dir)
        # print(f"Removed temp directory: {temp_dir}") # same here  (Removes temp audio folder)
        return result
    except Exception as e:
        return f"AssemblyAI Error: {str(e)}"


# Extracts key points and summary from the transcription using tokenization
def extract_key_points_and_summary(transcription):  # (Entry point 3)
    if not transcription or transcription.startswith("Error:"):
        return {"key_points": [], "summary": "No valid transcription provided."}

    sentences = sent_tokenize(transcription)
    if not sentences:
        return {"key_points": [], "summary": "No sentences found."}

    stop_words = set(stopwords.words('english'))
    words = [word.lower() for word in word_tokenize(transcription) if word.isalnum()]
    filtered_words = [word for word in words if word not in stop_words]

    word_freq = Counter(filtered_words)
    top_words = [word for word, freq in word_freq.most_common(10)]

    sentence_scores = {}
    for i, sentence in enumerate(sentences):
        score = sum(1 for word in word_tokenize(sentence.lower()) if word in top_words)
        sentence_scores[i] = score

    target_sentences = min(max(8, len(sentences) // 3), len(sentences))     # ensures at least 8 sentences are included
    top_sentences_idx = sorted(sentence_scores, key=sentence_scores.get, reverse=True)[:target_sentences]
    summary_sentences = [sentences[i] for i in sorted(top_sentences_idx)]
    summary = " ".join(summary_sentences)

    top_key_words = top_words[:5] 
    key_points = [s.strip() for s in sentences if any(word in s.lower() for word in top_key_words)]
    key_points = list(dict.fromkeys(key_points))[:10]       # Removes duplicates and limits to 10 key points

    summary_words = len(word_tokenize(summary))
    if summary_words < 150 and len(sentences) > len(summary_sentences):
        extra_sentences = [s for i, s in enumerate(sentences) if i not in top_sentences_idx][:2]
        summary = summary + " " + " ".join(extra_sentences)
    elif summary_words > 250 and len(summary_sentences) > 8:
        summary = " ".join(summary_sentences[:8])

    return {"key_points": key_points, "summary": summary}


# Refine the summary using Gemini AI model (entry point 4)
def refine_summary_with_gemini(summary, topics):

    prompt = f"""
    The following is a rough summary of a video. Please improve it to make it more structured, informative, and engaging while keeping it concise.
    Don't use any markdown just make a summary based conclusion in short upto 250 words based on the size of the Rough summary.
    **Topics Covered:** {topics}

    **Rough Summary:** {summary}

    Improve the summary while ensuring it properly covers all the topics & please don't oil it up, speak the actual thing because we will be using it for research purposes.

    If topics & summary not provided or says as error or API error or AssenblyAI error - or related error just type "Unexpected error occured" & nothing else.
    """
    response = model.generate_content(prompt)
    return response.text    # Returns the refined summary as text doc


# Generate short notes based on the refined summary and topics (entry point 5)
def short_notes(summary, topics):
    prompt = f"""
    The following is a summary of a video. Please analyze the summary and based on the topics covered provide notes on each related topic.
    if a topic is "pointers in cpp" then the notes should be like: what is pointers in cpp, how to use pointers in cpp, etc.
    Make sure to add a blank line between each topic.
    want all topics to be covered in the notes not only 4.
    **Topics Covered:** {topics}

    **Rough Summary:** {summary}

    While writing don't mention about the summary, just Topic: Ans (there must not be any mention of the summary in the notes).
    Make sure each topic covers at least 60 words.
    Apart from that mention 4-5 key points in bullet for each topic (eg. topic: Overview of PHP and its Comprehensive Coverage
    then the key points shoud be: characteristics of php & its Comprehensive Coverage).
    Remove markdown.
    If topics & summary not provided or says as error or API error or AssenblyAI error - or related error just type "Unexpected error occured" & nothing else.
    """
    output = model.generate_content(prompt)
    return output.text  # Returns the notes as text doc

# Extracts the video title from the URL (entry point 6)
def get_video_title(url):
    ydl_opts = {
        'quiet': True,
        'skip_download': True,
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info_dict = ydl.extract_info(url, download=False)
            video_title = info_dict.get('title', None)
            video_title = re.sub(r'[\\/*?:,"<>| #]', "_", video_title) # Removes all special characters from the title and replace them with '_'
            return video_title
    except yt_dlp.DownloadError as e:
        return f"Error: {str(e)}"
    
# Save the notes as a docx file (entry point 7)
def save_notes_as_docx(notes, filename="notes.docx"):
    filepath = os.path.join(OUTPUT_DIR, filename) # Save the file in the same directory as the script
    doc = Document()
    for line in notes.split('\n'):
        if line.startswith('*') and line.endswith('*'):
            doc.add_paragraph(line.strip('*')).bold = True
        else:
            doc.add_paragraph(line)
    doc.save(filepath)
    return filepath

# Save the data as a JSON file (entry point 8)
def save_to_json(data, filename="output.json"):
    filepath = os.path.join(OUTPUT_DIR, filename) # Save the file in the same directory as the script
    with open(filepath, 'w') as f:
        json.dump(data, f, indent=2)
    return filepath

# Check if the URL is a valid YouTube URL (entry point 0)
def is_valid_youtube_url(url):
    # Regular expression to check if the URL is a valid YouTube URL
    youtube_regex = re.compile(
        r'^(https?://)?(www\.)?(youtube|youtu|youtube-nocookie)\.(com|be)/.+$'
    )
    return youtube_regex.match(url) is not None

def check_video_exists(url):
    ydl_opts = {
        'quiet': True,
        'skip_download': True,
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.extract_info(url, download=False)
        return True
    except yt_dlp.DownloadError as e:
        logging.error(f"Video does not exist: {str(e)}")
        return False
    

# Convert text to audio using gTTS and pydub (entry point 9)
def text_to_audio(text, filename="temp_audio"):
    temp_filepath = os.path.join(TEMP_DIR, filename + '.mp3')
    # Create a temporary MP3 file using gTTS
    tts = gTTS(text=text, lang='en', slow=False)
    tts.save(temp_filepath)

    # Convert the MP3 file to a pydub AudioSegment
    audio = AudioSegment.from_mp3(temp_filepath)

    # Optionally increase playback speed (e.g., 1.2x for slightly faster speech)
    audio_speed = audio.speedup(playback_speed=1.2)  # Adjust speed as needed

    # Remove the temporary MP3 file
    os.remove(temp_filepath)

    return audio_speed




# Main entry point: Takes URL, returns JSON file (Enter point 1)
def process_video_url(url):
    
    """
    Checks if a URL exists by sending a HEAD request.
    Returns tuple: (exists: bool, message: str)
    """    
    # if not is_valid_youtube_url(url) or not check_video_exists(url):
    #     logging.warning(f"Invalid or non-existent YouTube URL: {url}")
    #     op = {
    #         "Transcript": "Invalid or non-existent YouTube URL.",
    #         "Summarization": "Invalid",
    #         "Notes": "Invalid",
    #         "NotesFile": "Invalid",
    #         "Audio": "Invalid"
    #     }
    #     return save_to_json(op, "output.json")
    
    try:
        # Use HEAD to avoid downloading content
        response = requests.head(url, allow_redirects=True, timeout=5)
        if response.status_code != 200:
            logging.warning(f"URL {url} returned status code: {response.status_code}")
            op = {
                "Transcript": f"URL returned status {response.status_code} (not found or inaccessible).",
                "Summarization": "Invalid",
                "Notes": "Invalid",
                "NotesFile": "Invalid",
                "Audio": "Invalid"
            }
            print(f"Output saved")
            return save_to_json(op, "output.json")
    except requests.RequestException as e:
        logging.error(f"Error checking URL {url}: {str(e)}")
        op = {
            "Transcript": f"Error checking URL: {str(e)}",
            "Summarization": "Invalid",
            "Notes": "Invalid",
            "NotesFile": "Invalid",
            "Audio": "Invalid"
        }
        print(f"Output saved")
        return save_to_json(op, "output.json")
  

    transcript = transcribe_video(url)
    print("Transcription done")
    # Ensure 'transcript' is a string and actually starts with 'Error'
    if isinstance(transcript, str) and transcript.startswith("Download Error:"):
        op = {
            "Transcript": transcript,
            "Summarization": "Invalid",
            "Notes": "Invalid",
            "NotesFile": "Invalid",
            "Audio": "Invalid"
        }
        print(f"Output saved")
        return save_to_json(op, "output.json")

    rough_result = extract_key_points_and_summary(transcript) # Extracts a rough summary and key points from the transcript
    # rough result is a dictionary with key_points and summary as keys
    
    refined_summary = refine_summary_with_gemini(rough_result["summary"], rough_result["key_points"]) # Refines the summary using Gemini AI model flask - 1.5
    print("Summary refined")
    notes = short_notes(refined_summary, rough_result["key_points"]) # Generates notes based on the refined summary & key points
    print("Notes generated")
    # video_title = get_video_title(url) # Extracts the video title from the URL
    # if video_title:
    #     filename = f"{video_title}.docx"
    # else:
    #     filename = "notes.docx"
    filename = "notes.docx"

    notes_file = save_notes_as_docx(notes,filename) # Saves the notes as a docx file

    audio=text_to_audio(refined_summary) # Converts the refined_summary to an audio file
    print("Audio generated")
    # Export the audio to a mp3 file
    audio_filepath = os.path.join(OUTPUT_DIR, 'Summary.mp3')
    audio.export(audio_filepath, format='mp3')


    final_result = {
        "Transcript": transcript,
        "Summarization": refined_summary,
        "Notes": notes,
        "NotesFile": notes_file,
        # "Audio": audio
    }    
    print(f"Output saved")
    return save_to_json(final_result) # Saves the final result as a JSON file & returns the JSON 



# Example usage (for testing, remove in production)
# if __name__ == "__main__":
#     video_url = "Enter URL"    
#     result = process_video_url(video_url)    

@app.route('/output/<path:filename>', methods=['GET'])
def download_file(filename):
    return send_from_directory(OUTPUT_DIR, filename)

@app.route('/process_url', methods=['POST'])
def process_url():
    data = request.json
    url = data.get('url')
    if not url:
        return jsonify({"error": "URL is required"}), 400

    result = process_video_url(url)
    return jsonify({"file_path": f"output/{os.path.basename(result)}"})

@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def serve_frontend(path):
    build_dir = os.path.join(os.path.dirname(__file__), 'dist')  # Use 'build' if CRA
    if path != "" and os.path.exists(os.path.join(build_dir, path)):
        return send_from_directory(build_dir, path)
    else:
        return send_from_directory(build_dir, 'index.html')
    
if __name__ == '__main__':
    app.run(debug=True)